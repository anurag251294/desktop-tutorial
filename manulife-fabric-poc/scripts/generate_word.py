"""Generate deal-sizing Word document."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    hs.font.name = "Calibri"


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()


# --- COVER PAGE ---
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Manulife Fabric Opportunity")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Deal Sizing \u2014 SWAG Estimate")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INTERNAL USE ONLY \u2014 DO NOT DISTRIBUTE")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run.bold = True

doc.add_paragraph()
for line in ["April 2026", "Prepared by: POC Delivery Team", "Classification: Internal"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# --- 1. OPPORTUNITY SUMMARY ---
doc.add_heading("1. Opportunity Summary", level=1)
doc.add_paragraph(
    "Manulife is evaluating Microsoft Fabric as the data foundation for a Copilot-style "
    "natural-language analytics experience across insurance and investment lines of business. "
    "The POC demonstrates structured and unstructured data integration through OneLake, a semantic "
    "model for trusted KPIs, and Fabric Data Agent as the governed query layer."
)

doc.add_heading("Customer Profile", level=2)
for item in [
    "~38,000 employees globally",
    "$900B+ assets under management",
    "Operations across Canada, US, and Asia",
    "Business lines: Group Benefits, Individual Insurance, Wealth & Asset Management, Retirement",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Initial Use Case", level=2)
doc.add_paragraph(
    "Claims and policy analytics with natural-language access for service representatives "
    "and business analysts, combining structured data insights with unstructured document context."
)

# --- 2. DEAL SUMMARY ---
doc.add_heading("2. Deal Summary", level=1)
add_table(
    ["Metric", "Low Estimate", "High Estimate", "Notes"],
    [
        ["POC Deal", "$150,000", "$200,000", "3-4 month engagement"],
        ["Year 1 Production", "$1,000,000", "$2,000,000", "Single BU production deployment"],
        ["Annual Run Rate (Yr 2+)", "$600,000", "$1,500,000", "Licensing + managed services"],
        ["3-Year TCV (Base Case)", "$3,000,000", "$6,000,000", "2-3 BUs, moderate expansion"],
        ["3-Year TCV (Upside)", "$6,000,000", "$12,000,000", "Enterprise-wide, Copilot at scale"],
    ],
    [5, 4, 4, 6],
)

# --- 3. POC PHASE ---
doc.add_heading("3. POC / Pilot Phase (3-4 Months)", level=1)
add_table(
    ["Component", "Monthly", "Total (4 mo)", "Notes"],
    [
        ["Fabric Capacity (F64)", "$7,000", "$28,000", "Minimum viable for POC"],
        ["Azure AI Search (Basic)", "$1,000", "$4,000", "Single index, low volume"],
        ["Azure OpenAI (S0)", "$2,000", "$8,000", "Embeddings + GPT-4o"],
        ["Power BI Pro (10 seats)", "$1,000", "$4,000", "Dev and testing"],
        ["Azure Subtotal", "$11,000", "$44,000", ""],
        ["Consulting", "", "$75K-$150K", "Architecture & implementation"],
        ["POC TOTAL", "", "$150K-$200K", "Target landing deal"],
    ],
    [5, 3, 3.5, 6],
)

# --- 4. YEAR 1 PRODUCTION ---
doc.add_heading("4. Year 1 \u2014 Production Deployment", level=1)
add_table(
    ["Component", "Low Annual", "High Annual", "Notes"],
    [
        ["Fabric Capacity (F128-F256)", "$200,000", "$500,000", "Production workloads"],
        ["Power BI Premium", "$60,000", "$250,000", "Capacity vs seat model"],
        ["Azure OpenAI", "$100,000", "$250,000", "Production query volume"],
        ["Azure AI Search (Std)", "$50,000", "$120,000", "Vector search, replicas"],
        ["Supporting Azure", "$10,000", "$30,000", "KV, Storage, Monitor"],
        ["Implementation", "$300,000", "$800,000", "Hardening, security, rollout"],
        ["YEAR 1 TOTAL", "$720,000", "$1,950,000", "Target: $1M-$2M"],
    ],
    [5, 3.5, 3.5, 6],
)

# --- 5. STEADY STATE ---
doc.add_heading("5. Steady-State Annual Run Rate (Year 2+)", level=1)
add_table(
    ["Component", "Low Annual", "High Annual", "Notes"],
    [
        ["Fabric Capacity", "$250,000", "$600,000", "Grows with data volume"],
        ["Power BI Licensing", "$60,000", "$250,000", "Stable unless seats expand"],
        ["Azure AI Services", "$150,000", "$350,000", "Scales with queries"],
        ["Managed Services", "$150,000", "$400,000", "L2/L3, tuning, enhancements"],
        ["ANNUAL RUN RATE", "$610,000", "$1,600,000", ""],
    ],
    [5, 3.5, 3.5, 6],
)

# --- 6. EXPANSION POTENTIAL ---
doc.add_heading("6. Expansion Potential", level=1)
doc.add_paragraph(
    "The initial use case is a wedge into a much larger opportunity. "
    "Manulife has multiple business units that could adopt the same pattern."
)

doc.add_heading("3-Year Scenarios", level=2)
add_table(
    ["Scenario", "Assumptions", "Low TCV", "High TCV"],
    [
        ["Conservative", "Single BU, contained scope", "$1.5M", "$3M"],
        ["Base Case", "2-3 BUs, moderate expansion", "$3M", "$6M"],
        ["Upside", "Enterprise-wide, Copilot at scale", "$6M", "$12M+"],
    ],
    [3.5, 7, 3, 3],
)

doc.add_heading("Expansion Vectors", level=2)
add_table(
    ["Vector", "Description", "Multiplier"],
    [
        ["Additional BUs", "Group Benefits, Indiv. Insurance, Wealth, Retirement", "2x-4x"],
        ["Data Domains", "Actuarial, risk, compliance, finance", "1.5x-2x"],
        ["Copilot Seats", "Pilot (50-100) to enterprise (5,000+)", "Major uplift"],
        ["Data Volume", "More sources, real-time streams", "1.3x-2x"],
        ["Geography", "US and Asia operations", "1.5x-2x"],
        ["Advanced AI", "Agentic workflows, auto triage", "Incremental"],
    ],
    [4, 8, 3.5],
)

# --- 7. RISKS ---
doc.add_page_break()
doc.add_heading("7. Risks and Assumptions", level=1)

doc.add_heading("Risks to Deal Size", level=2)
add_table(
    ["Risk", "Impact", "Likelihood", "Mitigation"],
    [
        ["Existing ELA covers Fabric", "Reduces net-new licensing", "Medium", "Validate licensing early"],
        ["Data Agent stays preview", "Delays production", "Medium", "Pivot to PBI Copilot"],
        ["Competing platform", "Limits scope", "Low-Med", "Position as complementary"],
        ["Budget freeze", "Extends cycle", "Low", "Align to budget cycles"],
        ["Internal build pref.", "Reduces consulting", "Medium", "Demonstrate POC velocity"],
    ],
    [4, 4.5, 3, 5],
)

doc.add_heading("Key Assumptions", level=2)
for a in [
    "Manulife does not already have significant Fabric / Power BI Premium licensing",
    "Fabric Data Agent reaches GA in a compatible timeframe",
    "POC successfully demonstrates value to business stakeholders",
    "Partner-delivered consulting is the preferred delivery model",
    "No competing platform commitment blocks Fabric adoption",
]:
    doc.add_paragraph(a, style="List Bullet")

doc.add_heading("Upside Drivers", level=2)
for a in [
    "Microsoft co-sell / co-invest interest (Manulife is a strategic account)",
    "Regulatory mandate to modernize data platform",
    "Competitive pressure from peers adopting AI analytics",
    "Executive sponsor with transformation mandate",
]:
    doc.add_paragraph(a, style="List Bullet")

# --- 8. COMPETITIVE ---
doc.add_heading("8. Competitive Landscape", level=1)
add_table(
    ["Competitor", "Threat", "Our Positioning"],
    [
        ["Databricks + Unity Catalog", "Med-High", "Fabric: unified platform, no BI/AI stitching"],
        ["Snowflake + Cortex", "Medium", "Power BI + Copilot integration differentiated"],
        ["AWS (Bedrock + Redshift)", "Low-Med", "Manulife likely has Azure affinity"],
        ["Internal / Custom", "Medium", "POC demonstrates speed-to-value"],
    ],
    [5, 3, 8.5],
)

# --- 9. NEXT STEPS ---
doc.add_heading("9. Recommended Next Steps", level=1)
add_table(
    ["#", "Action", "Owner", "Timeline"],
    [
        ["1", "Validate current Microsoft licensing posture", "Account Team", "Week 1"],
        ["2", "Identify executive sponsor and budget holder", "Account Team", "Week 1-2"],
        ["3", "Present POC results and reference architecture", "Delivery Team", "Week 2"],
        ["4", "Scope production pilot (single BU)", "Joint", "Week 3-4"],
        ["5", "Develop SOW for production phase", "Delivery Team", "Week 4-5"],
        ["6", "Engage Microsoft co-sell / FastTrack", "Account Team", "Week 2-3"],
        ["7", "Go/no-go for enterprise expansion", "Joint", "Month 3-4"],
    ],
    [1.5, 8, 3.5, 3],
)

# --- 10. BOTTOM LINE ---
doc.add_heading("10. Bottom Line", level=1)
p = doc.add_paragraph()
run = p.add_run(
    "Land with a $150K-$200K POC, convert to a $1M-$2M Year 1 production deal, "
    "and expand to $3M-$6M over 3 years across multiple business units. "
    "Microsoft co-sell alignment and Fabric Data Agent GA timing are the two biggest swing factors."
)
run.bold = True

filepath = "C:/Users/anuragdhuria/OneDrive - Microsoft/Documents/GitHub/desktop-tutorial/manulife-fabric-poc/docs/deal-sizing-swag.docx"
doc.save(filepath)
print(f"Word doc saved: {filepath}")

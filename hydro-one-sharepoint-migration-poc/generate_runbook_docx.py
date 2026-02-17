"""
Generate Hydro One SharePoint Migration Runbook as Word Document
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ============================================================
# STYLES
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title page style
title_style = doc.styles['Title']
title_style.font.size = Pt(28)
title_style.font.color.rgb = RGBColor(0, 51, 102)

# Heading styles
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.color.rgb = RGBColor(0, 51, 102)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    # Data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)

def add_note(doc, text, prefix="NOTE"):
    p = doc.add_paragraph()
    run = p.add_run(f"{prefix}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)
    run = p.add_run(text)
    run.font.italic = True

def add_checkbox(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f"\u2610  {text}")
    run.font.size = Pt(11)

# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Hydro One\nSharePoint to Azure Migration")
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0, 51, 102)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Step-by-Step Deployment Runbook")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f"Version 1.1 | {datetime.date.today().strftime('%B %Y')}")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

add_table(doc,
    ["Field", "Value"],
    [
        ["Project", "Hydro One SharePoint to Azure Data Lake Migration"],
        ["Document Type", "Deployment Runbook"],
        ["Author", "Microsoft Azure Data Engineering Team"],
        ["Classification", "Confidential"],
        ["Version", "1.1"],
        ["Last Updated", datetime.date.today().strftime("%B %d, %Y")],
        ["Data Volume", "~25 TB"],
        ["Target Platform", "Azure Data Lake Storage Gen2"],
    ])

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Critical Prerequisites (MUST Complete First)",
    "2. Step 1: Azure Subscription Setup",
    "3. Step 2: Register Resource Providers",
    "4. Step 3: Create Resource Group",
    "5. Step 4: Create ADLS Gen2 Storage Account",
    "6. Step 5: Create Azure SQL Server and Database",
    "7. Step 6: Create Azure Key Vault",
    "8. Step 7: Register Azure AD Application",
    "9. Step 8: Grant Admin Consent (CRITICAL)",
    "10. Step 9: Deploy Azure Data Factory",
    "11. Step 10: Deploy ADF Pipelines",
    "12. Step 11: Initialize SQL Database",
    "13. Step 12: Grant ADF Permissions",
    "14. Step 13: Verify All Connections",
    "15. Step 14: Populate Control Table",
    "16. Step 15: Run Pilot Migration",
    "17. Step 16: Monitor Migration",
    "18. Step 17: Run Validation Pipeline",
    "19. Step 18: Full-Scale Migration",
    "20. Step 19: Enable Incremental Sync",
    "21. Troubleshooting Guide",
    "22. Monitoring SQL Queries",
    "23. Rollback Procedures",
    "24. Test Results from POC",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# SECTION 1: CRITICAL PREREQUISITES
# ============================================================
doc.add_heading('1. Critical Prerequisites', level=1)
add_note(doc, "These items MUST be completed before starting the deployment. Without them, the migration WILL NOT work.", "BLOCKER")
doc.add_paragraph()

doc.add_heading('1.1 Required Azure Roles', level=2)
add_table(doc,
    ["Role", "Who Needs It", "Purpose"],
    [
        ["Contributor", "Deployment engineer", "Create Azure resources"],
        ["Application Administrator", "Azure AD admin", "Register app and add API permissions"],
        ["Global Administrator", "IT admin", "Grant admin consent for SharePoint API"],
        ["SharePoint Administrator", "SharePoint admin", "Access to site collections"],
    ])

doc.add_paragraph()
doc.add_heading('1.2 Required Software', level=2)
add_table(doc,
    ["Tool", "Version", "Download"],
    [
        ["Azure CLI", "2.50+", "https://aka.ms/installazurecli"],
        ["PowerShell", "7.0+", "https://aka.ms/powershell"],
        ["Az PowerShell Module", "10.0+", "Install-Module -Name Az"],
        ["SQL Server Mgmt Studio", "19+", "https://aka.ms/ssms"],
        ["Git", "2.40+", "https://git-scm.com"],
    ])

doc.add_paragraph()
doc.add_heading('1.3 Information Needed Before Starting', level=2)
add_checkbox(doc, "Azure Subscription ID")
add_checkbox(doc, "Azure AD Tenant ID")
add_checkbox(doc, "Admin user UPN (e.g., admin@hydroone.onmicrosoft.com)")
add_checkbox(doc, "Admin user Object ID (from Azure AD)")
add_checkbox(doc, "List of SharePoint site collections to migrate")
add_checkbox(doc, "Estimated file counts and sizes per library")
add_checkbox(doc, "Preferred Azure region (recommended: Canada Central)")
add_checkbox(doc, "Environment name (dev / test / prod)")

doc.add_paragraph()
doc.add_heading('1.4 Network Requirements', level=2)
doc.add_paragraph("Ensure the following outbound connections are allowed:")
add_checkbox(doc, "HTTPS (443) to *.sharepoint.com")
add_checkbox(doc, "HTTPS (443) to *.azure.com and *.azure.net")
add_checkbox(doc, "HTTPS (443) to login.microsoftonline.com")
add_checkbox(doc, "No firewall blocking Azure Data Factory managed runtime IPs")

doc.add_page_break()

# ============================================================
# STEP 1: AZURE SUBSCRIPTION SETUP
# ============================================================
doc.add_heading('2. Step 1: Azure Subscription Setup', level=1)

doc.add_heading('2.1 Login to Azure', level=2)
doc.add_paragraph("Open a terminal (PowerShell or Bash) and run:")
add_code_block(doc, 'az login --tenant "<your-tenant-id>"')
doc.add_paragraph("This will open a browser window. Sign in with your Azure admin credentials.")

doc.add_heading('2.2 Set the Target Subscription', level=2)
add_code_block(doc, 'az account set --subscription "<your-subscription-id>"')

doc.add_heading('2.3 Verify Your Context', level=2)
add_code_block(doc, 'az account show')
doc.add_paragraph("Confirm the output shows the correct subscription name and tenant ID.")

doc.add_heading('2.4 Record Your Details', level=2)
add_table(doc,
    ["Item", "Your Value"],
    [
        ["Subscription ID", ""],
        ["Tenant ID", ""],
        ["Admin UPN", ""],
        ["Admin Object ID", ""],
        ["Region", "canadacentral"],
        ["Environment", ""],
    ])

doc.add_page_break()

# ============================================================
# STEP 2: REGISTER RESOURCE PROVIDERS
# ============================================================
doc.add_heading('3. Step 2: Register Resource Providers', level=1)
doc.add_paragraph("Azure resource providers must be registered before creating resources. Run each command:")

add_code_block(doc, """az provider register --namespace Microsoft.DataFactory
az provider register --namespace Microsoft.Storage
az provider register --namespace Microsoft.Sql
az provider register --namespace Microsoft.KeyVault""")

doc.add_paragraph("Wait for registration to complete (may take 1-2 minutes). Verify each one:")
add_code_block(doc, """az provider show --namespace Microsoft.DataFactory --query "registrationState" -o tsv
az provider show --namespace Microsoft.Storage --query "registrationState" -o tsv
az provider show --namespace Microsoft.Sql --query "registrationState" -o tsv
az provider show --namespace Microsoft.KeyVault --query "registrationState" -o tsv""")

doc.add_paragraph('All four should return "Registered".')
add_checkbox(doc, "Microsoft.DataFactory = Registered")
add_checkbox(doc, "Microsoft.Storage = Registered")
add_checkbox(doc, "Microsoft.Sql = Registered")
add_checkbox(doc, "Microsoft.KeyVault = Registered")

doc.add_page_break()

# ============================================================
# STEP 3: CREATE RESOURCE GROUP
# ============================================================
doc.add_heading('4. Step 3: Create Resource Group', level=1)
doc.add_paragraph("Create a resource group to contain all migration resources:")
add_code_block(doc, """az group create \\
    --name "rg-hydroone-migration-{env}" \\
    --location "canadacentral" \\
    --tags "Project=HydroOneMigration" "Environment={env}" "Owner=Microsoft" """)
add_note(doc, 'Replace {env} with your environment name (dev, test, or prod) throughout this document.')
add_checkbox(doc, "Resource group created successfully")

doc.add_page_break()

# ============================================================
# STEP 4: CREATE ADLS GEN2 STORAGE
# ============================================================
doc.add_heading('5. Step 4: Create ADLS Gen2 Storage Account', level=1)

doc.add_heading('5.1 Create the Storage Account', level=2)
add_code_block(doc, """az storage account create \\
    --name "sthydroonemig{env}" \\
    --resource-group "rg-hydroone-migration-{env}" \\
    --location "canadacentral" \\
    --sku "Standard_LRS" \\
    --kind "StorageV2" \\
    --enable-hierarchical-namespace true \\
    --min-tls-version "TLS1_2" """)
add_note(doc, "The --enable-hierarchical-namespace true flag enables ADLS Gen2 (Data Lake). This CANNOT be changed after creation.")

doc.add_heading('5.2 Create Storage Containers', level=2)
add_code_block(doc, """# Primary migration container
az storage container create \\
    --name "sharepoint-migration" \\
    --account-name "sthydroonemig{env}" \\
    --auth-mode login

# Metadata container
az storage container create \\
    --name "migration-metadata" \\
    --account-name "sthydroonemig{env}" \\
    --auth-mode login""")

add_checkbox(doc, "Storage account created with hierarchical namespace enabled")
add_checkbox(doc, "sharepoint-migration container created")
add_checkbox(doc, "migration-metadata container created")

doc.add_page_break()

# ============================================================
# STEP 5: CREATE SQL
# ============================================================
doc.add_heading('6. Step 5: Create Azure SQL Server and Database', level=1)

doc.add_heading('6.1 Create SQL Server', level=2)
doc.add_paragraph("Create an Azure SQL Server with Azure AD authentication:")
add_code_block(doc, """az sql server create \\
    --name "sql-hydroone-migration-{env}" \\
    --resource-group "rg-hydroone-migration-{env}" \\
    --location "canadacentral" \\
    --enable-ad-only-auth \\
    --external-admin-principal-type "User" \\
    --external-admin-name "<your-admin-upn>" \\
    --external-admin-sid "<your-admin-object-id>" """)
add_note(doc, "Replace <your-admin-upn> with your Azure AD email and <your-admin-object-id> with your Object ID from Azure AD > Users.")

doc.add_heading('6.2 Create Database', level=2)
add_code_block(doc, """az sql db create \\
    --name "MigrationControl" \\
    --server "sql-hydroone-migration-{env}" \\
    --resource-group "rg-hydroone-migration-{env}" \\
    --service-objective "S1" \\
    --backup-storage-redundancy "Local" """)

doc.add_heading('6.3 Configure Firewall Rules', level=2)
add_code_block(doc, """# Allow Azure services (required for ADF)
az sql server firewall-rule create \\
    --name "AllowAzureServices" \\
    --server "sql-hydroone-migration-{env}" \\
    --resource-group "rg-hydroone-migration-{env}" \\
    --start-ip-address 0.0.0.0 \\
    --end-ip-address 0.0.0.0

# Allow your client IP (for running SQL scripts)
az sql server firewall-rule create \\
    --name "AllowClientIP" \\
    --server "sql-hydroone-migration-{env}" \\
    --resource-group "rg-hydroone-migration-{env}" \\
    --start-ip-address "<your-public-ip>" \\
    --end-ip-address "<your-public-ip>" """)
add_note(doc, 'Find your public IP at https://whatismyip.com')

add_checkbox(doc, "SQL Server created with Azure AD admin")
add_checkbox(doc, "MigrationControl database created")
add_checkbox(doc, "Firewall rules configured")

doc.add_page_break()

# ============================================================
# STEP 6: CREATE KEY VAULT
# ============================================================
doc.add_heading('7. Step 6: Create Azure Key Vault', level=1)
add_code_block(doc, """az keyvault create \\
    --name "kv-hydroone-mig-{env}" \\
    --resource-group "rg-hydroone-migration-{env}" \\
    --location "canadacentral" \\
    --enable-rbac-authorization false \\
    --sku "standard" """)
add_note(doc, "We use access policies (not RBAC) for Key Vault to simplify permission management.")
add_checkbox(doc, "Key Vault created")

doc.add_page_break()

# ============================================================
# STEP 7: REGISTER AZURE AD APP
# ============================================================
doc.add_heading('8. Step 7: Register Azure AD Application', level=1)

doc.add_heading('8.1 Create App Registration', level=2)
add_code_block(doc, """az ad app create \\
    --display-name "HydroOne-SPO-Migration" \\
    --sign-in-audience "AzureADMyOrg" """)
doc.add_paragraph('Save the "appId" from the output — you will need it in the next steps.')

doc.add_heading('8.2 Create Service Principal', level=2)
add_code_block(doc, 'az ad sp create --id "<app-id>"')

doc.add_heading('8.3 Generate Client Secret', level=2)
add_code_block(doc, """az ad app credential reset \\
    --id "<app-id>" \\
    --years 1""")
add_note(doc, "IMPORTANT: Copy the 'password' value immediately. It cannot be retrieved later.", "WARNING")

doc.add_heading('8.4 Add SharePoint API Permissions', level=2)
doc.add_paragraph("Add the required SharePoint and Graph API permissions:")
add_code_block(doc, """# SharePoint - Sites.FullControl.All (Application)
az ad app permission add \\
    --id "<app-id>" \\
    --api "00000003-0000-0ff1-ce00-000000000000" \\
    --api-permissions "678536fe-1083-478a-9c59-b99265e6b0d3=Role"

# SharePoint - Sites.ReadWrite.All (Application)
az ad app permission add \\
    --id "<app-id>" \\
    --api "00000003-0000-0ff1-ce00-000000000000" \\
    --api-permissions "fbcd29d2-fcca-4405-aded-518d457caae4=Role"

# Microsoft Graph - Sites.ReadWrite.All (Application)
az ad app permission add \\
    --id "<app-id>" \\
    --api "00000003-0000-0000-c000-000000000000" \\
    --api-permissions "9492366f-7969-46a4-8d15-ed1a20078fff=Role" """)

doc.add_heading('8.5 Store Credentials in Key Vault', level=2)
add_code_block(doc, """az keyvault secret set \\
    --vault-name "kv-hydroone-mig-{env}" \\
    --name "sharepoint-client-secret" \\
    --value "<client-secret-value>"

az keyvault secret set \\
    --vault-name "kv-hydroone-mig-{env}" \\
    --name "sharepoint-client-id" \\
    --value "<app-id>"

az keyvault secret set \\
    --vault-name "kv-hydroone-mig-{env}" \\
    --name "tenant-id" \\
    --value "<tenant-id>" """)

doc.add_heading('8.6 Record Your App Details', level=2)
add_table(doc,
    ["Item", "Your Value"],
    [
        ["App (Client) ID", ""],
        ["Service Principal Object ID", ""],
        ["Client Secret", "(stored in Key Vault)"],
        ["Tenant ID", ""],
    ])

add_checkbox(doc, "App registration created")
add_checkbox(doc, "Service principal created")
add_checkbox(doc, "Client secret generated and saved")
add_checkbox(doc, "API permissions added")
add_checkbox(doc, "Secrets stored in Key Vault")

doc.add_page_break()

# ============================================================
# STEP 8: ADMIN CONSENT (CRITICAL)
# ============================================================
doc.add_heading('9. Step 8: Grant Admin Consent (CRITICAL)', level=1)
p = doc.add_paragraph()
run = p.add_run("THIS IS THE MOST CRITICAL STEP. Without admin consent, the migration WILL NOT WORK.")
run.bold = True
run.font.color.rgb = RGBColor(180, 0, 0)
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph("A Global Administrator must perform these steps:")
doc.add_paragraph()

doc.add_heading('9.1 Steps in Azure Portal', level=2)
steps = [
    'Sign in to https://portal.azure.com as a Global Administrator',
    'Navigate to Azure Active Directory in the left menu',
    'Click "App registrations" in the left menu',
    'Click the "All applications" tab',
    'Search for and select "HydroOne-SPO-Migration"',
    'Click "API permissions" in the left menu',
    'You will see the permissions listed with "Not granted" status',
    'Click the "Grant admin consent for Hydro One" button (blue button at top)',
    'A confirmation dialog appears — click "Yes"',
    'Wait for the operation to complete (10-30 seconds)',
    'Verify ALL permissions now show a green checkmark with "Granted for Hydro One"',
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph(f"Step {i}: ", style='List Number')
    p.add_run(step)

doc.add_paragraph()
doc.add_heading('9.2 Required Permissions (must all show "Granted")', level=2)
add_table(doc,
    ["API", "Permission", "Type", "Status After Consent"],
    [
        ["SharePoint", "Sites.FullControl.All", "Application", "Granted ✓"],
        ["SharePoint", "Sites.ReadWrite.All", "Application", "Granted ✓"],
        ["Microsoft Graph", "Sites.ReadWrite.All", "Application", "Granted ✓"],
    ])

doc.add_paragraph()
doc.add_heading('9.3 Alternative: Azure CLI', level=2)
add_code_block(doc, 'az ad app permission admin-consent --id "<app-id>"')
add_note(doc, "This command also requires Global Administrator role.")

doc.add_paragraph()
doc.add_heading('9.4 Verification', level=2)
doc.add_paragraph("After granting consent, verify by running:")
add_code_block(doc, """az ad app permission list-grants \\
    --id "<app-id>" \\
    --query "[].{resource:resourceDisplayName, scope:scope}" -o table""")

add_checkbox(doc, "Admin consent granted by Global Administrator")
add_checkbox(doc, "All permissions show 'Granted' status in portal")

doc.add_page_break()

# ============================================================
# STEP 9: DEPLOY ADF
# ============================================================
doc.add_heading('10. Step 9: Deploy Azure Data Factory', level=1)

doc.add_heading('10.1 Clone the Repository', level=2)
add_code_block(doc, """git clone https://github.com/anurag251294/desktop-tutorial.git
cd desktop-tutorial/hydro-one-sharepoint-migration-poc""")

doc.add_heading('10.2 Deploy ADF ARM Template', level=2)
doc.add_paragraph("This deploys the ADF instance with all linked services and datasets:")
add_code_block(doc, """az deployment group create \\
    --resource-group "rg-hydroone-migration-{env}" \\
    --template-file "adf-templates/arm-template.json" \\
    --parameters \\
        factoryName="adf-hydroone-migration-{env}" \\
        location="canadacentral" \\
        sharePointTenantUrl="https://hydroone.sharepoint.com" \\
        servicePrincipalId="<app-id>" \\
        tenantId="<tenant-id>" \\
        keyVaultName="kv-hydroone-mig-{env}" \\
        storageAccountName="sthydroonemig{env}" \\
        sqlServerName="sql-hydroone-migration-{env}" \\
        sqlDatabaseName="MigrationControl" """)

doc.add_paragraph("This creates:")
add_table(doc,
    ["Resource", "Count", "Names"],
    [
        ["ADF Instance", "1", "adf-hydroone-migration-{env}"],
        ["Linked Services", "5", "LS_AzureKeyVault, LS_SharePointOnline_REST, LS_SharePointOnline_HTTP, LS_ADLS_Gen2, LS_AzureSqlDatabase"],
        ["Datasets", "5", "DS_SharePoint_Binary_HTTP, DS_ADLS_Binary_Sink, DS_ADLS_Parquet_Metadata, DS_SQL_MigrationControl, DS_SQL_AuditLog"],
    ])

doc.add_heading('10.3 Get ADF Managed Identity', level=2)
doc.add_paragraph("Record the ADF managed identity — you'll need it for permission grants:")
add_code_block(doc, """az datafactory show \\
    --name "adf-hydroone-migration-{env}" \\
    --resource-group "rg-hydroone-migration-{env}" \\
    --query "identity.principalId" -o tsv""")
doc.add_paragraph("Save this value: ____________________________")

add_checkbox(doc, "ADF deployed with managed identity")
add_checkbox(doc, "ADF managed identity ID recorded")

doc.add_page_break()

# ============================================================
# STEP 10: DEPLOY PIPELINES
# ============================================================
doc.add_heading('11. Step 10: Deploy ADF Pipelines', level=1)
doc.add_paragraph("Deploy each pipeline in order (some have dependencies):")

pipelines = [
    ("1", "PL_Process_Subfolder", "pipeline-subfolder", "None"),
    ("2", "PL_Migrate_Single_Library", "pipeline-single-library", "PL_Process_Subfolder"),
    ("3", "PL_Validation", "pipeline-validation", "None"),
    ("4", "PL_Master_Migration_Orchestrator", "pipeline-master", "PL_Migrate_Single_Library"),
    ("5", "PL_Incremental_Sync", "pipeline-incremental", "None"),
]

for num, name, deploy_name, dep in pipelines:
    doc.add_heading(f'Pipeline {num}: {name}', level=2)
    if dep != "None":
        add_note(doc, f"Depends on {dep} — deploy that one first.")
    add_code_block(doc, f"""az deployment group create \\
    --resource-group "rg-hydroone-migration-{{env}}" \\
    --template-file "adf-templates/pipelines/{name}.json" \\
    --parameters factoryName="adf-hydroone-migration-{{env}}" \\
    --name "{deploy_name}" """)
    add_checkbox(doc, f"{name} deployed successfully")

doc.add_heading('Verify All Pipelines', level=2)
add_code_block(doc, """az rest --method GET \\
    --url "https://management.azure.com/subscriptions/{sub-id}/resourceGroups/rg-hydroone-migration-{env}/providers/Microsoft.DataFactory/factories/adf-hydroone-migration-{env}/pipelines?api-version=2018-06-01" \\
    --query "value[].name" -o table""")

doc.add_paragraph("Expected output: 5 pipelines listed.")

doc.add_page_break()

# ============================================================
# STEP 11: INITIALIZE SQL
# ============================================================
doc.add_heading('12. Step 11: Initialize SQL Database', level=1)

doc.add_heading('12.1 Connect to the Database', level=2)
doc.add_paragraph("Option A: Azure Portal")
steps = [
    'Go to Azure Portal > SQL databases > MigrationControl',
    'Click "Query editor (preview)" in the left menu',
    'Sign in with Azure Active Directory authentication',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph()
doc.add_paragraph("Option B: SQL Server Management Studio (SSMS)")
steps = [
    'Open SSMS',
    'Server name: sql-hydroone-migration-{env}.database.windows.net',
    'Authentication: Azure Active Directory - Universal with MFA',
    'Connect',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('12.2 Run Script 1: Control Tables', level=2)
doc.add_paragraph('Open and execute the entire contents of: sql/create_control_table.sql')
doc.add_paragraph("This creates:")
add_table(doc,
    ["Object", "Type", "Purpose"],
    [
        ["MigrationControl", "Table", "Tracks all libraries to migrate"],
        ["IncrementalWatermark", "Table", "High watermark for delta sync"],
        ["BatchLog", "Table", "Batch execution tracking"],
        ["SyncLog", "Table", "Incremental sync tracking"],
        ["usp_UpdateMigrationStatus", "Stored Procedure", "Update library status"],
        ["usp_LogBatchStart", "Stored Procedure", "Log batch start"],
        ["usp_LogBatchComplete", "Stored Procedure", "Log batch completion"],
        ["usp_UpdateValidationStatus", "Stored Procedure", "Update validation status"],
        ["usp_LogValidationResult", "Stored Procedure", "Log validation results"],
        ["usp_UpdateWatermark", "Stored Procedure", "Update incremental sync watermark"],
        ["usp_LogSyncRun", "Stored Procedure", "Log sync run completion"],
    ])

doc.add_heading('12.3 Run Script 2: Audit Log Tables', level=2)
doc.add_paragraph('Open and execute: sql/create_audit_log_table.sql')
doc.add_paragraph("This creates:")
add_table(doc,
    ["Object", "Type", "Purpose"],
    [
        ["MigrationAuditLog", "Table", "Per-file migration audit trail"],
        ["ValidationLog", "Table", "Validation results"],
        ["usp_LogFileAudit", "Stored Procedure", "Log individual file migration"],
        ["usp_BulkLogFileAudit", "Stored Procedure", "Bulk insert audit records"],
    ])

doc.add_heading('12.4 Run Script 3: Monitoring Queries', level=2)
doc.add_paragraph('Open and execute: sql/monitoring_queries.sql')

doc.add_heading('12.5 Verify Tables', level=2)
add_code_block(doc, """SELECT TABLE_NAME
FROM INFORMATION_SCHEMA.TABLES
WHERE TABLE_SCHEMA = 'dbo'
ORDER BY TABLE_NAME;""")
doc.add_paragraph("Expected tables: BatchLog, IncrementalWatermark, MigrationAuditLog, MigrationControl, SyncLog, ValidationLog")

add_checkbox(doc, "Control tables created (Script 1)")
add_checkbox(doc, "Audit tables created (Script 2)")
add_checkbox(doc, "Monitoring queries created (Script 3)")
add_checkbox(doc, "All 6 tables verified")

doc.add_page_break()

# ============================================================
# STEP 12: GRANT ADF PERMISSIONS
# ============================================================
doc.add_heading('13. Step 12: Grant ADF Managed Identity Permissions', level=1)
doc.add_paragraph("The ADF managed identity needs access to Key Vault, Storage, and SQL.")

doc.add_heading('13.1 Grant Key Vault Access', level=2)
add_code_block(doc, """az keyvault set-policy \\
    --name "kv-hydroone-mig-{env}" \\
    --object-id "<adf-managed-identity-id>" \\
    --secret-permissions get list""")

doc.add_heading('13.2 Grant Storage Access', level=2)
add_code_block(doc, """az role assignment create \\
    --role "Storage Blob Data Contributor" \\
    --assignee-object-id "<adf-managed-identity-id>" \\
    --assignee-principal-type "ServicePrincipal" \\
    --scope "/subscriptions/{sub-id}/resourceGroups/rg-hydroone-migration-{env}/providers/Microsoft.Storage/storageAccounts/sthydroonemig{env}" """)

doc.add_heading('13.3 Grant SQL Access', level=2)
doc.add_paragraph("In SQL Query Editor, run:")
add_code_block(doc, """CREATE USER [adf-hydroone-migration-{env}] FROM EXTERNAL PROVIDER;
ALTER ROLE db_datareader ADD MEMBER [adf-hydroone-migration-{env}];
ALTER ROLE db_datawriter ADD MEMBER [adf-hydroone-migration-{env}];
GRANT EXECUTE ON SCHEMA::dbo TO [adf-hydroone-migration-{env}];""")

add_checkbox(doc, "Key Vault access policy set for ADF")
add_checkbox(doc, "Storage Blob Data Contributor role assigned to ADF")
add_checkbox(doc, "SQL user created for ADF managed identity")

doc.add_page_break()

# ============================================================
# STEP 13: VERIFY CONNECTIONS
# ============================================================
doc.add_heading('14. Step 13: Verify All Connections', level=1)
doc.add_paragraph("In Azure Data Factory Studio (https://adf.azure.com):")

steps = [
    'Open ADF Studio and select your factory',
    'Go to Manage (toolbox icon) > Linked services',
    'Click "Test connection" on each linked service:',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

add_table(doc,
    ["Linked Service", "Expected Result", "If It Fails"],
    [
        ["LS_AzureKeyVault", "Connection successful", "Check ADF Key Vault access policy"],
        ["LS_ADLS_Gen2", "Connection successful", "Check Storage Blob Data Contributor role"],
        ["LS_AzureSqlDatabase", "Connection successful", "Check SQL user was created for ADF"],
        ["LS_SharePointOnline_REST", "Connection successful", "Check admin consent was granted"],
        ["LS_SharePointOnline_HTTP", "Connection successful", "Check SharePoint URL is correct"],
    ])

add_checkbox(doc, "All 5 linked services show 'Connection successful'")

doc.add_page_break()

# ============================================================
# STEP 14: POPULATE CONTROL TABLE
# ============================================================
doc.add_heading('15. Step 14: Populate Control Table', level=1)

doc.add_heading('15.1 Manual Entry (for testing)', level=2)
doc.add_paragraph("Insert a test library into the control table:")
add_code_block(doc, """INSERT INTO dbo.MigrationControl (
    SiteUrl, LibraryName, SiteTitle, LibraryTitle,
    FileCount, TotalSizeBytes,
    Status, Priority, EnableIncrementalSync
)
VALUES (
    '/sites/YourTestSite',
    'Documents',
    'YourTestSite',
    'Documents',
    100,            -- estimated file count
    1073741824,     -- estimated size (1 GB)
    'Pending',
    1,
    1
);""")

doc.add_heading('15.2 Automated Population (for full migration)', level=2)
add_code_block(doc, """.\\scripts\\Populate-ControlTable.ps1 `
    -SharePointTenantUrl "https://hydroone.sharepoint.com" `
    -SqlServerName "sql-hydroone-migration-{env}" `
    -SqlDatabaseName "MigrationControl" `
    -UseInteractiveAuth""")

doc.add_heading('15.3 Verify Control Table', level=2)
add_code_block(doc, """SELECT SiteUrl, LibraryName, FileCount,
    CAST(TotalSizeBytes / 1073741824.0 AS DECIMAL(10,2)) AS SizeGB,
    Status, Priority
FROM dbo.MigrationControl
ORDER BY Priority, TotalSizeBytes DESC;""")

add_checkbox(doc, "Control table populated with library records")

doc.add_page_break()

# ============================================================
# STEP 15: PILOT MIGRATION
# ============================================================
doc.add_heading('16. Step 15: Run Pilot Migration', level=1)
add_note(doc, "Start with a small library (< 1 GB, < 100 files) to validate everything works.", "RECOMMENDATION")

doc.add_heading('16.1 Find Smallest Library', level=2)
add_code_block(doc, """SELECT TOP 1 Id, SiteUrl, LibraryName, FileCount,
    CAST(TotalSizeBytes / 1048576.0 AS DECIMAL(10,2)) AS SizeMB
FROM dbo.MigrationControl
WHERE Status = 'Pending'
ORDER BY TotalSizeBytes ASC;""")

doc.add_heading('16.2 Trigger Pilot Run', level=2)
steps = [
    'Open ADF Studio > Author > Pipelines > SharePoint Migration',
    'Select PL_Master_Migration_Orchestrator',
    'Click "Debug" button',
    'Set parameters: BatchSize=1, ParallelLibraries=1, MaxRetries=3, TargetContainerName=sharepoint-migration',
    'Click "OK" to start the pipeline',
    'Monitor progress in the output panel',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('16.3 Monitor Pilot', level=2)
doc.add_paragraph("In ADF Monitor tab, click on the running pipeline to see activity progress.")
doc.add_paragraph("In SQL, check the migration status:")
add_code_block(doc, """-- Check library status
SELECT * FROM dbo.MigrationControl WHERE Status = 'InProgress';

-- Check file-level progress
SELECT MigrationStatus, COUNT(*) AS Files,
    SUM(FileSizeBytes) / 1048576.0 AS TotalMB
FROM dbo.MigrationAuditLog
WHERE BatchId = (SELECT TOP 1 BatchId FROM dbo.BatchLog ORDER BY StartTime DESC)
GROUP BY MigrationStatus;""")

doc.add_heading('16.4 Verify Pilot in ADLS', level=2)
steps = [
    'Go to Azure Portal > Storage Account > Containers > sharepoint-migration',
    'Navigate to the site/library folder structure',
    'Verify files exist and sizes look correct',
    'Open a few files to confirm they are not corrupted',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

add_checkbox(doc, "Pilot migration completed successfully")
add_checkbox(doc, "Files verified in ADLS")

doc.add_page_break()

# ============================================================
# STEP 16: MONITOR
# ============================================================
doc.add_heading('17. Step 16: Monitor Migration', level=1)

doc.add_heading('17.1 PowerShell Monitoring', level=2)
add_code_block(doc, """.\\scripts\\Monitor-Migration.ps1 `
    -ResourceGroupName "rg-hydroone-migration-{env}" `
    -DataFactoryName "adf-hydroone-migration-{env}" `
    -SqlServerName "sql-hydroone-migration-{env}" `
    -SqlDatabaseName "MigrationControl" `
    -ContinuousMonitor""")

doc.add_heading('17.2 Key SQL Monitoring Queries', level=2)
doc.add_paragraph("Overall Progress:")
add_code_block(doc, """SELECT Status, COUNT(*) AS Libraries,
    CAST(SUM(TotalSizeBytes) / 1099511627776.0 AS DECIMAL(10,2)) AS TotalTB
FROM dbo.MigrationControl
GROUP BY Status;""")

doc.add_paragraph("Daily Throughput:")
add_code_block(doc, """SELECT CAST([Timestamp] AS DATE) AS [Date],
    COUNT(*) AS FilesCopied,
    CAST(SUM(FileSizeBytes) / 1073741824.0 AS DECIMAL(10,2)) AS GB_Copied
FROM dbo.MigrationAuditLog
WHERE MigrationStatus = 'Success'
GROUP BY CAST([Timestamp] AS DATE)
ORDER BY [Date] DESC;""")

doc.add_paragraph("Throttling Detection:")
add_code_block(doc, """SELECT CAST([Timestamp] AS DATE) AS [Date],
    DATEPART(HOUR, [Timestamp]) AS [Hour],
    COUNT(*) AS ThrottleCount
FROM dbo.MigrationAuditLog
WHERE ErrorCode = '429'
GROUP BY CAST([Timestamp] AS DATE), DATEPART(HOUR, [Timestamp])
ORDER BY [Date] DESC, [Hour];""")

doc.add_page_break()

# ============================================================
# STEP 17: VALIDATION
# ============================================================
doc.add_heading('18. Step 17: Run Validation Pipeline', level=1)
doc.add_paragraph("After migration batches complete, run validation to verify data integrity:")

doc.add_heading('18.1 Trigger Validation', level=2)
steps = [
    'In ADF Studio, go to Author > Pipelines > PL_Validation',
    'Click "Debug"',
    'Set ValidateAll = true',
    'Click "OK"',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('18.2 Check Validation Results', level=2)
add_code_block(doc, """SELECT SiteUrl, LibraryName, FileCount AS Expected,
    MigratedFileCount AS Migrated, FailedFileCount AS Failed,
    ValidationStatus, DiscrepancyDetails
FROM dbo.MigrationControl
WHERE Status = 'Completed'
ORDER BY ValidationStatus;""")

doc.add_heading('18.3 Validation Sign-Off Checklist', level=2)
add_checkbox(doc, "All libraries show ValidationStatus = 'Validated'")
add_checkbox(doc, "File count matches within 1% tolerance")
add_checkbox(doc, "Total size matches within 5% tolerance")
add_checkbox(doc, "No critical files missing")
add_checkbox(doc, "Spot-check samples pass verification")
add_checkbox(doc, "Business stakeholder sign-off obtained")

doc.add_page_break()

# ============================================================
# STEP 18: FULL SCALE
# ============================================================
doc.add_heading('19. Step 18: Full-Scale Migration', level=1)

doc.add_heading('19.1 Recommended Batch Schedule', level=2)
add_table(doc,
    ["Phase", "Days", "Time", "BatchSize", "Parallelism", "Target"],
    [
        ["Week 1", "Mon-Fri", "8PM-6AM", "5", "2", "~2 TB"],
        ["Week 2", "Mon-Sun", "8PM-6AM / All day", "10", "4", "~5 TB"],
        ["Week 3-4", "Mon-Sun", "8PM-6AM / All day", "15", "4", "~8 TB"],
        ["Week 5-6", "Mon-Sun", "8PM-6AM / All day", "20", "4", "~10 TB"],
    ])

doc.add_heading('19.2 Trigger Full Migration', level=2)
add_code_block(doc, """Invoke-AzDataFactoryV2Pipeline `
    -ResourceGroupName "rg-hydroone-migration-{env}" `
    -DataFactoryName "adf-hydroone-migration-{env}" `
    -PipelineName "PL_Master_Migration_Orchestrator" `
    -Parameter @{
        BatchSize = 10
        ParallelLibraries = 4
        MaxRetries = 3
        TargetContainerName = "sharepoint-migration"
    }""")

doc.add_page_break()

# ============================================================
# STEP 19: INCREMENTAL SYNC
# ============================================================
doc.add_heading('20. Step 19: Enable Incremental Sync', level=1)
doc.add_paragraph("After initial migration is validated, enable incremental sync for ongoing changes:")

doc.add_heading('20.1 Enable Sync for Validated Libraries', level=2)
add_code_block(doc, """UPDATE dbo.MigrationControl
SET EnableIncrementalSync = 1
WHERE Status = 'Completed'
AND ValidationStatus = 'Validated';""")

doc.add_heading('20.2 Initialize Watermark', level=2)
add_code_block(doc, """INSERT INTO dbo.IncrementalWatermark (SiteUrl, LibraryName, LastModifiedDate, LastSyncTime)
SELECT SiteUrl, LibraryName, EndTime, GETUTCDATE()
FROM dbo.MigrationControl
WHERE EnableIncrementalSync = 1;""")

doc.add_heading('20.3 Start Sync Trigger', level=2)
add_code_block(doc, """Start-AzDataFactoryV2Trigger `
    -ResourceGroupName "rg-hydroone-migration-{env}" `
    -DataFactoryName "adf-hydroone-migration-{env}" `
    -Name "TR_TumblingWindow_IncrementalSync" """)

add_checkbox(doc, "Incremental sync enabled and trigger started")

doc.add_page_break()

# ============================================================
# TROUBLESHOOTING
# ============================================================
doc.add_heading('21. Troubleshooting Guide', level=1)

add_table(doc,
    ["Error", "Cause", "Resolution"],
    [
        ["HTTP 401 Unauthorized", "Token expired or invalid credentials", "Regenerate client secret, update Key Vault, verify admin consent"],
        ["HTTP 403 Forbidden", "Insufficient permissions", "Verify admin consent granted, check site-level permissions"],
        ["HTTP 404 Not Found", "File deleted after enumeration", "Expected — log and skip, file was removed at source"],
        ["HTTP 429 Throttled", "SharePoint API rate limit exceeded", "Reduce parallelism, wait, contact Microsoft for limit increase"],
        ["HTTP 503 Service Unavailable", "SharePoint temporarily down", "Automatic retry, check service health dashboard"],
        ["SQL connection failed", "Firewall or permissions issue", "Check firewall rules, verify ADF SQL user created"],
        ["Key Vault access denied", "Missing access policy", "Run az keyvault set-policy for ADF identity"],
        ["Storage access denied", "Missing role assignment", "Assign Storage Blob Data Contributor to ADF identity"],
        ["ARM template [dbo] error", "ARM interprets [ as expression", "Use [[dbo] in ARM templates (escapes to literal [)"],
        ["ForEach batchCount error", "Expression used for batchCount", "Use static integer 1-50 (not a parameter expression)"],
    ])

doc.add_page_break()

# ============================================================
# ROLLBACK
# ============================================================
doc.add_heading('23. Rollback Procedures', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Key Principle: SharePoint data is READ-ONLY during migration. No changes are made to the source. Full rollback = continue using SharePoint as before.")
run.bold = True

doc.add_heading('23.1 Rollback Steps', level=2)
steps = [
    'Stop all ADF triggers',
    'Cancel any running pipelines in ADF Monitor',
    'Notify stakeholders',
    'Export audit logs for analysis',
    'If ADLS data is corrupted: delete container contents and re-migrate',
    'Reset control table: UPDATE dbo.MigrationControl SET Status = \'Pending\', RetryCount = 0 WHERE Status = \'Failed\'',
    'Address root cause before retry',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_page_break()

# ============================================================
# TEST RESULTS
# ============================================================
doc.add_heading('24. Test Results from POC', level=1)
doc.add_paragraph("The following tests were performed in the Microsoft MCAPS lab environment (February 2026):")

doc.add_heading('24.1 PL_Master_Migration_Orchestrator', level=2)
doc.add_paragraph("Run ID: 7bb1ca75-5180-4087-9870-6dea48e667a6 — Status: SUCCEEDED")
add_table(doc,
    ["Activity", "Status", "Duration"],
    [
        ["Set_BatchId", "Succeeded", "276ms"],
        ["Log_BatchStart", "Succeeded", "3,808ms"],
        ["Lookup_PendingLibraries", "Succeeded", "8,019ms"],
        ["If_NoLibrariesToProcess", "Succeeded", "978ms"],
        ["Filter_BatchSize", "Succeeded", "319ms"],
        ["Log_NoWork", "Succeeded", "254ms"],
        ["ForEach_Library", "Succeeded", "756ms"],
        ["Log_BatchComplete", "Succeeded", "3,071ms"],
    ])

doc.add_heading('24.2 PL_Validation', level=2)
doc.add_paragraph("Run ID: 92158504-c822-46a3-9a4f-2cbd780986f6 — Status: SUCCEEDED")
add_table(doc,
    ["Activity", "Status", "Duration"],
    [
        ["Lookup_CompletedLibraries", "Succeeded", "26,107ms"],
        ["ForEach_ValidateLibrary", "Succeeded", "40,775ms"],
        ["Lookup_DestinationFileCount", "Succeeded", "24,090ms"],
        ["Compare_And_Log", "Succeeded", "3,717ms"],
        ["If_Discrepancy", "Succeeded", "5,365ms"],
        ["Mark_Validated", "Succeeded", "2,817ms"],
        ["Generate_ValidationReport", "Succeeded", "13,051ms"],
        ["Set_ValidationSummary", "Succeeded", "276ms"],
    ])

doc.add_heading('24.3 Verified Connectivity', level=2)
add_table(doc,
    ["Connection", "Status"],
    [
        ["ADF → Azure SQL (read/write)", "Working"],
        ["ADF → Azure SQL (stored procedures)", "Working"],
        ["ADF → Key Vault (secrets)", "Working"],
        ["ADF → ADLS Gen2 (storage)", "Role Assigned"],
        ["ADF → SharePoint Online", "Requires admin consent in target tenant"],
    ])

# ============================================================
# SAVE
# ============================================================
output_path = r"C:\Users\anuragdhuria\OneDrive - Microsoft\Documents\GitHub\desktop-tutorial\hydro-one-sharepoint-migration-poc\docs\Hydro_One_SharePoint_Migration_Runbook.docx"
doc.save(output_path)
print(f"Word document saved to: {output_path}")

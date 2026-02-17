"""
Generate Word documents for each pipeline and master architecture doc
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

BASE = r"C:\Users\anuragdhuria\OneDrive - Microsoft\Documents\GitHub\desktop-tutorial\hydro-one-sharepoint-migration-poc\docs"

def make_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    for i in range(1, 4):
        doc.styles[f'Heading {i}'].font.color.rgb = RGBColor(0, 51, 102)
    return doc

def add_title_page(doc, title, subtitle):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0, 51, 102)
    r.bold = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Microsoft Azure Data Engineering Team | {datetime.date.today().strftime('%B %Y')}")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_page_break()

def tbl(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(30, 30, 30)

def note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("Note: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0, 100, 0)
    p.add_run(text)

def flow_diagram(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0, 51, 102)

# ================================================================
# 1. MASTER ARCHITECTURE DOCUMENT
# ================================================================
def gen_architecture():
    doc = make_doc()
    add_title_page(doc,
        "Hydro One SharePoint\nto Azure Migration",
        "Solution Architecture Document")

    # 1. Executive Summary
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "This document describes the architecture for migrating approximately 25 TB of documents "
        "from Hydro One's SharePoint Online environment to Azure Data Lake Storage Gen2 (ADLS Gen2). "
        "The solution uses Azure Data Factory (ADF) as the orchestration engine, with metadata-driven "
        "pipelines that provide automated, resumable, and auditable migration capabilities.")
    doc.add_paragraph(
        "The migration is designed to run during off-peak hours to minimize impact on SharePoint "
        "performance, with built-in throttle detection and exponential backoff to respect Microsoft 365 "
        "service limits.")

    doc.add_heading('1.1 Key Design Decisions', level=2)
    tbl(doc,
        ["Decision", "Choice", "Rationale"],
        [
            ["Orchestration Engine", "Azure Data Factory", "Native Azure service, managed infrastructure, built-in SharePoint connectors, visual monitoring"],
            ["Authentication", "Managed Identity + Service Principal", "MI for Azure services (no credentials), SP for SharePoint OAuth2"],
            ["Destination Storage", "ADLS Gen2 (hierarchical namespace)", "Folder-like structure preserves SharePoint hierarchy, optimized for analytics"],
            ["Control Database", "Azure SQL Database", "Relational tracking, transactional updates, rich querying for monitoring"],
            ["Secret Management", "Azure Key Vault", "Centralized, audited, rotatable credentials"],
            ["File Transfer", "Binary copy via HTTP", "Preserves file integrity, no format conversion, streaming (no staging)"],
            ["Error Handling", "Retry with exponential backoff", "Handles transient failures, throttling (HTTP 429), and service outages"],
            ["Validation", "SQL-based count/size comparison", "No additional SharePoint API calls needed for validation"],
        ])

    doc.add_page_break()

    # 2. Architecture Overview
    doc.add_heading('2. Architecture Overview', level=1)

    doc.add_heading('2.1 High-Level Architecture', level=2)
    flow_diagram(doc, """
    +----------------------+                    +-------------------+
    |                      |   REST API (HTTPS)  |                   |
    |  SharePoint Online   | -----------------> |  Azure Data       |
    |  (25 TB Source)      | <----------------- |  Factory          |
    |                      |   OAuth2 Bearer     |  (Orchestration)  |
    +----------------------+                    +--------+----------+
                                                         |
                            +----------------------------+----------------------------+
                            |                            |                            |
                   +--------v---------+         +--------v---------+        +---------v--------+
                   |                  |         |                  |        |                  |
                   |  ADLS Gen2       |         |  Azure SQL DB    |        |  Azure Key Vault |
                   |  (Destination)   |         |  (Control/Audit) |        |  (Secrets)       |
                   |  25 TB storage   |         |  Migration state |        |  Client secret   |
                   |                  |         |                  |        |                  |
                   +------------------+         +------------------+        +------------------+
    """)

    doc.add_heading('2.2 Data Flow', level=2)
    doc.add_paragraph("The migration follows a pull-based model:")
    steps = [
        "ADF reads the MigrationControl table to identify pending libraries",
        "For each library, ADF calls SharePoint REST API to enumerate files in the document library",
        "For each file, ADF downloads the binary content via HTTP GET",
        "The binary stream is written directly to ADLS Gen2 (no intermediate staging)",
        "Each file copy result (success/failure) is logged to the MigrationAuditLog table",
        "After all files in a library are processed, the library status is updated to Completed or Failed",
        "The validation pipeline compares expected vs actual file counts to verify data integrity",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {s}")

    doc.add_page_break()

    # 3. Component Details
    doc.add_heading('3. Component Details', level=1)

    doc.add_heading('3.1 Azure Data Factory', level=2)
    doc.add_paragraph(
        "ADF serves as the central orchestration engine. It uses a system-assigned managed identity "
        "for passwordless authentication to Azure services (SQL, ADLS, Key Vault). For SharePoint access, "
        "it uses an OAuth2 service principal whose credentials are stored in Key Vault.")
    tbl(doc,
        ["Property", "Value"],
        [
            ["SKU", "Pay-as-you-go (no dedicated runtime needed)"],
            ["Identity", "System-assigned Managed Identity"],
            ["Linked Services", "5 (Key Vault, SharePoint REST, SharePoint HTTP, ADLS, SQL)"],
            ["Datasets", "5 (SharePoint Binary, ADLS Binary Sink, ADLS Parquet, SQL Control, SQL Audit)"],
            ["Pipelines", "5 (Orchestrator, Single Library, Subfolder, Validation, Incremental Sync)"],
            ["Concurrency", "4 libraries in parallel, 10 files per library"],
        ])

    doc.add_heading('3.2 ADLS Gen2 Storage', level=2)
    doc.add_paragraph("Files are stored in a hierarchical folder structure that mirrors the SharePoint site/library layout:")
    flow_diagram(doc, """
    sharepoint-migration/                    <-- Container
    +-- SiteName1/                           <-- SharePoint Site
    |   +-- Documents/                       <-- Document Library
    |   |   +-- Reports/                     <-- Subfolder
    |   |   |   +-- Q1_Report.pdf            <-- File
    |   |   |   +-- Q2_Report.pdf
    |   |   +-- Policies/
    |   |       +-- Safety_Policy.docx
    |   +-- Shared Documents/
    |       +-- ...
    +-- SiteName2/
        +-- ...
    """)

    doc.add_heading('3.3 Azure SQL Database', level=2)
    doc.add_paragraph("Stores all migration state, audit trails, and validation results.")
    tbl(doc,
        ["Table", "Purpose", "Key Columns"],
        [
            ["MigrationControl", "Tracks each library to migrate", "SiteUrl, LibraryName, Status, FileCount, MigratedFileCount"],
            ["MigrationAuditLog", "Per-file migration audit trail", "FileName, SourcePath, DestinationPath, MigrationStatus, FileSizeBytes"],
            ["BatchLog", "Batch execution tracking", "BatchId, Status, StartTime, EndTime"],
            ["IncrementalWatermark", "Delta sync high watermark", "SiteUrl, LibraryName, LastModifiedDate"],
            ["ValidationLog", "Validation results", "ValidationType, ExpectedValue, ActualValue, IsMatch"],
            ["SyncLog", "Incremental sync tracking", "SyncType, LibrariesProcessed, CompletionTime"],
        ])

    doc.add_heading('3.4 Azure Key Vault', level=2)
    doc.add_paragraph("Stores sensitive credentials used by the migration pipelines:")
    tbl(doc,
        ["Secret Name", "Contains", "Used By"],
        [
            ["sharepoint-client-secret", "OAuth2 client secret for SharePoint access", "LS_SharePointOnline_REST"],
            ["sharepoint-client-id", "Azure AD Application (client) ID", "LS_SharePointOnline_REST"],
            ["tenant-id", "Azure AD Tenant ID", "LS_SharePointOnline_REST"],
        ])

    doc.add_page_break()

    # 4. Pipeline Architecture
    doc.add_heading('4. Pipeline Architecture', level=1)

    doc.add_heading('4.1 Pipeline Hierarchy', level=2)
    flow_diagram(doc, """
    PL_Master_Migration_Orchestrator              <-- Entry point
        |
        +-- Reads MigrationControl table (Pending/Failed libraries)
        +-- Filters to BatchSize
        +-- ForEach Library (parallel, batch=4)
             |
             +-- PL_Migrate_Single_Library        <-- Per-library migration
                  |
                  +-- Updates status to InProgress
                  +-- Enumerates root files via SharePoint REST API
                  +-- ForEach Root File (parallel, batch=10)
                  |    +-- Binary copy: SharePoint -> ADLS
                  |    +-- Log success/failure to audit table
                  |    +-- Detect and handle HTTP 429 throttling
                  |
                  +-- Enumerates subfolders via SharePoint REST API
                  +-- ForEach Subfolder
                       |
                       +-- PL_Process_Subfolder   <-- Per-folder processing
                            +-- Enumerate files in subfolder
                            +-- ForEach File (parallel, batch=10)
                                 +-- Binary copy: SharePoint -> ADLS
                                 +-- Log success/failure

    PL_Validation                                 <-- Post-migration
        +-- Reads completed libraries from control table
        +-- Compares expected vs actual file counts
        +-- Marks as Validated or Discrepancy

    PL_Incremental_Sync                           <-- Ongoing
        +-- Reads validated libraries with sync enabled
        +-- Queries SharePoint for files modified since last watermark
        +-- Copies only changed files
        +-- Updates watermark
    """)

    doc.add_heading('4.2 Pipeline Summary', level=2)
    tbl(doc,
        ["Pipeline", "Purpose", "Trigger", "Parallelism"],
        [
            ["PL_Master_Migration_Orchestrator", "Top-level batch orchestrator", "Manual / Scheduled", "4 libraries"],
            ["PL_Migrate_Single_Library", "Migrate all files from one library", "Called by Orchestrator", "10 files"],
            ["PL_Process_Subfolder", "Process files in one subfolder", "Called by Single Library", "10 files"],
            ["PL_Validation", "Validate migration completeness", "Manual (post-migration)", "5 libraries"],
            ["PL_Incremental_Sync", "Ongoing delta sync", "Tumbling window (6 hrs)", "4 libraries"],
        ])

    doc.add_page_break()

    # 5. Authentication & Security
    doc.add_heading('5. Authentication & Security', level=1)

    doc.add_heading('5.1 Authentication Flows', level=2)
    tbl(doc,
        ["Connection", "Authentication Method", "Credential Storage"],
        [
            ["ADF -> ADLS Gen2", "Managed Identity", "None (Azure AD token)"],
            ["ADF -> Azure SQL", "Managed Identity", "None (Azure AD token)"],
            ["ADF -> Key Vault", "Managed Identity", "None (Azure AD token)"],
            ["ADF -> SharePoint (REST)", "OAuth2 Client Credentials", "Key Vault (client secret)"],
            ["ADF -> SharePoint (HTTP)", "MSI / Bearer Token", "Pipeline-level token injection"],
        ])

    doc.add_heading('5.2 OAuth2 Flow for SharePoint', level=2)
    flow_diagram(doc, """
    1. ADF Managed Identity authenticates to Key Vault
    2. Retrieves client-id and client-secret from Key Vault
    3. Sends POST to Azure AD token endpoint:
       POST https://accounts.accesscontrol.windows.net/{tenant}/tokens/OAuth/2
       Body: grant_type=client_credentials&client_id={id}&client_secret={secret}&resource=00000003-0000-0ff1-ce00-000000000000/{tenant}.sharepoint.com@{tenant-id}
    4. Receives Bearer access token
    5. Uses token in Authorization header for SharePoint REST API calls
    """)

    doc.add_heading('5.3 Security Principles', level=2)
    principles = [
        "No credentials stored in code or pipeline definitions — all secrets in Key Vault",
        "Managed Identity used wherever possible (eliminates credential management)",
        "SharePoint access is READ-ONLY — no modifications to source data",
        "SQL Database uses Azure AD only authentication (no SQL passwords)",
        "Key Vault uses access policies with least-privilege (get/list only for ADF)",
        "ADLS uses role-based access (Storage Blob Data Contributor for ADF)",
        "All data transfer over HTTPS (TLS 1.2+)",
    ]
    for p_text in principles:
        doc.add_paragraph(p_text, style='List Bullet')

    doc.add_page_break()

    # 6. Throttling & Performance
    doc.add_heading('6. Throttling & Performance Management', level=1)

    doc.add_heading('6.1 SharePoint Online Limits', level=2)
    tbl(doc,
        ["Limit Type", "Threshold", "ADF Response"],
        [
            ["Requests/minute", "~600", "HTTP 429 triggers wait + retry"],
            ["Concurrent connections", "~10-15", "ForEach batchCount limits concurrency"],
            ["File download bandwidth", "Variable", "Streaming (no buffering)"],
            ["Retry-After header", "1-300 seconds", "Pipeline waits for specified duration"],
        ])

    doc.add_heading('6.2 Throttle Detection Strategy', level=2)
    flow_diagram(doc, """
    File Copy Activity
         |
         +-- Success? --> Log to audit table, continue
         |
         +-- HTTP 429? --> Wait (ThrottleWaitSeconds parameter, default 120s)
         |                      |
         |                      +-- Retry (up to 3 attempts)
         |
         +-- HTTP 401/403? --> Log error, skip file
         |
         +-- HTTP 404? --> Log as expected (file deleted at source), skip
         |
         +-- Timeout? --> Retry with extended timeout
         |
         +-- Other error? --> Log error, increment retry count
    """)

    doc.add_heading('6.3 Recommended Scheduling', level=2)
    tbl(doc,
        ["Time Window (EST)", "Activity", "Parallelism", "Rationale"],
        [
            ["6 AM - 6 PM", "Minimal / None", "1-2 libraries", "Business hours, avoid impacting users"],
            ["6 PM - 10 PM", "Moderate", "4 libraries", "Reduced user activity"],
            ["10 PM - 6 AM", "Maximum", "4 libraries", "Off-peak, lowest throttling risk"],
            ["Weekends", "Maximum", "4 libraries", "Minimal user activity"],
        ])

    doc.add_page_break()

    # 7. Monitoring & Observability
    doc.add_heading('7. Monitoring & Observability', level=1)

    doc.add_heading('7.1 Monitoring Layers', level=2)
    tbl(doc,
        ["Layer", "Tool", "What It Shows"],
        [
            ["Pipeline Level", "ADF Monitor", "Pipeline runs, activity status, durations, errors"],
            ["File Level", "SQL MigrationAuditLog", "Per-file status, sizes, error codes, timestamps"],
            ["Library Level", "SQL MigrationControl", "Library status, file counts, validation status"],
            ["Batch Level", "SQL BatchLog", "Batch start/end, libraries completed/failed"],
            ["Throttling", "SQL AuditLog (ErrorCode=429)", "Throttle frequency by hour/day"],
            ["Throughput", "SQL AuditLog aggregation", "GB/hour, files/hour"],
        ])

    doc.add_heading('7.2 Key Metrics', level=2)
    metrics = [
        "Migration Progress: % of libraries completed, TB migrated vs total",
        "Throughput: GB/hour, files/hour (should sustain 50-100 GB/hour off-peak)",
        "Error Rate: % of files failed, by error type (429, 401, 404, timeout)",
        "Throttle Frequency: 429 errors per hour (target: < 10/hour)",
        "Retry Rate: files requiring retry (target: < 5%)",
        "Validation Pass Rate: % of libraries passing validation (target: 100%)",
    ]
    for m in metrics:
        doc.add_paragraph(m, style='List Bullet')

    doc.add_page_break()

    # 8. Disaster Recovery
    doc.add_heading('8. Disaster Recovery & Rollback', level=1)
    doc.add_paragraph(
        "The migration is designed to be safe, resumable, and fully reversible:")

    tbl(doc,
        ["Principle", "Implementation"],
        [
            ["Source is read-only", "SharePoint data is never modified during migration"],
            ["Idempotent operations", "Re-running a pipeline overwrites destination (same result)"],
            ["Resumable", "Failed libraries can be reset to Pending and re-processed"],
            ["Audited", "Every file copy is logged with source path, status, and error details"],
            ["Rollback = no action", "Since source is unchanged, rollback means continuing to use SharePoint"],
        ])

    doc.add_paragraph()
    doc.add_heading('8.1 Recovery Scenarios', level=2)
    tbl(doc,
        ["Scenario", "Recovery Action", "Data Loss Risk"],
        [
            ["Pipeline fails mid-batch", "Reset failed libraries to Pending, re-run", "None (source unchanged)"],
            ["Credential expires", "Regenerate secret, update Key Vault, re-run", "None"],
            ["ADLS data corrupted", "Delete container, reset control table, re-migrate", "None (source is authoritative)"],
            ["SQL database lost", "Restore from backup, or re-create tables and re-migrate", "Audit history lost"],
            ["Full rollback", "Stop triggers, continue using SharePoint", "None"],
        ])

    doc.add_page_break()

    # 9. Cost Estimation
    doc.add_heading('9. Cost Estimation', level=1)
    tbl(doc,
        ["Resource", "SKU/Tier", "Monthly Cost (CAD)", "Notes"],
        [
            ["ADLS Gen2 (25 TB, Hot)", "Standard LRS", "$650-800", "Includes storage + transactions"],
            ["Azure SQL Database", "S1 (20 DTU)", "$40", "Sufficient for control/audit tables"],
            ["Azure Data Factory", "Pay-as-you-go", "$65-260", "Depends on activity volume"],
            ["Azure Key Vault", "Standard", "$7", "Minimal secret operations"],
            ["Total", "", "$762-1,107/month", ""],
        ])
    note(doc, "Costs are estimates based on Canada Central pricing. Actual costs may vary.")

    # 10. Deployment Environments
    doc.add_heading('10. Deployment Environments', level=1)
    tbl(doc,
        ["Environment", "Purpose", "Region", "Naming Pattern"],
        [
            ["dev", "Development and initial testing", "Canada Central", "rg-hydroone-migration-dev"],
            ["test", "Pre-production validation", "Canada Central", "rg-hydroone-migration-test"],
            ["prod", "Production migration", "Canada Central", "rg-hydroone-migration-prod"],
        ])

    doc.save(f"{BASE}\\Architecture_Overview.docx")
    print("Created: Architecture_Overview.docx")

# ================================================================
# 2. PL_MASTER_MIGRATION_ORCHESTRATOR
# ================================================================
def gen_pl_master():
    doc = make_doc()
    add_title_page(doc, "PL_Master_Migration\nOrchestrator", "Pipeline Technical Documentation")

    doc.add_heading('1. Overview', level=1)
    tbl(doc,
        ["Property", "Value"],
        [
            ["Pipeline Name", "PL_Master_Migration_Orchestrator"],
            ["ARM Template", "adf-templates/pipelines/PL_Master_Migration_Orchestrator.json"],
            ["Purpose", "Top-level orchestrator — reads pending libraries from the control table and processes them in parallel batches"],
            ["Trigger", "Manual, scheduled trigger, or tumbling window"],
            ["Child Pipeline", "PL_Migrate_Single_Library"],
            ["Test Status", "PASSED (Run ID: 7bb1ca75, all 8 activities succeeded)"],
        ])

    doc.add_heading('2. Parameters', level=1)
    tbl(doc,
        ["Parameter", "Type", "Default", "Description"],
        [
            ["BatchSize", "int", "10", "Maximum number of libraries to process in this run"],
            ["ParallelLibraries", "int", "4", "Reserved for future use (ForEach uses static batchCount=4)"],
            ["MaxRetries", "int", "3", "Maximum retry attempts per library before skipping"],
            ["TargetContainerName", "string", "sharepoint-migration", "ADLS Gen2 container name"],
        ])
    note(doc, "ADF ForEach batchCount must be a static integer (1-50). The parallelism is hardcoded to 4 in the pipeline definition, not driven by the ParallelLibraries parameter.")

    doc.add_heading('3. Variables', level=1)
    tbl(doc,
        ["Variable", "Type", "Description"],
        [
            ["BatchId", "String", "Auto-generated batch identifier (format: BATCH-yyyyMMdd-HHmmss)"],
            ["NoWorkMessage", "String", "Message set when no pending libraries are found"],
        ])

    doc.add_heading('4. Activity Flow', level=1)
    flow_diagram(doc, """
    +-------------+      +----------------+      +-------------------------+
    | Set_BatchId | ---> | Log_BatchStart | ---> | Lookup_PendingLibraries |
    +-------------+      +----------------+      +------------+------------+
                                                              |
                                            +-----------------+------------------+
                                            |                                    |
                                  +---------v-----------+            +-----------v-----------+
                                  | If_NoLibrariesToProcess |         | Filter_BatchSize       |
                                  +---------+-----------+            +-----------+-----------+
                                            |                                    |
                                   (if empty)                           +--------v--------+
                                  +---------v-----------+              | ForEach_Library  |
                                  | Log_NoWork          |              | (parallel, 4)    |
                                  +---------------------+              +--------+---------+
                                                                                |
                                                                    +-----------v-----------+
                                                                    | Execute_Migrate       |
                                                                    | SingleLibrary         |
                                                                    +-----------+-----------+
                                                                                |
                                                                    +-----------v-----------+
                                                                    | Log_BatchComplete     |
                                                                    +-----------------------+
    """)

    doc.add_heading('5. Activity Details', level=1)

    doc.add_heading('5.1 Set_BatchId', level=2)
    tbl(doc, ["Property", "Value"], [
        ["Type", "SetVariable"],
        ["Expression", "@concat('BATCH-', formatDateTime(utcNow(), 'yyyyMMdd-HHmmss'))"],
        ["Purpose", "Generates a unique batch identifier for tracking this run"],
        ["Example Output", "BATCH-20260217-200000"],
    ])

    doc.add_heading('5.2 Log_BatchStart', level=2)
    tbl(doc, ["Property", "Value"], [
        ["Type", "SqlServerStoredProcedure"],
        ["Stored Procedure", "dbo.usp_LogBatchStart"],
        ["Parameters", "BatchId, PipelineRunId, StartTime"],
        ["Purpose", "Creates a record in BatchLog table to track this batch execution"],
    ])

    doc.add_heading('5.3 Lookup_PendingLibraries', level=2)
    tbl(doc, ["Property", "Value"], [
        ["Type", "Lookup (firstRowOnly: false)"],
        ["Dataset", "DS_SQL_MigrationControl"],
        ["Timeout", "10 minutes"],
        ["Retry", "3 attempts, 30s interval"],
    ])
    doc.add_paragraph("SQL Query:")
    code(doc, """SELECT Id, SiteUrl, LibraryName, FileCount, TotalSizeBytes, RetryCount
FROM dbo.MigrationControl
WHERE Status IN ('Pending', 'Failed')
AND (RetryCount < @{pipeline().parameters.MaxRetries} OR RetryCount IS NULL)
ORDER BY TotalSizeBytes ASC""")
    note(doc, "Libraries are ordered by size (smallest first) to achieve early wins and validate the pipeline quickly.")

    doc.add_heading('5.4 Filter_BatchSize', level=2)
    tbl(doc, ["Property", "Value"], [
        ["Type", "Filter"],
        ["Purpose", "Limits the number of libraries to the BatchSize parameter"],
        ["Expression", "@lessOrEquals(indexOf(...), sub(pipeline().parameters.BatchSize, 1))"],
    ])

    doc.add_heading('5.5 ForEach_Library', level=2)
    tbl(doc, ["Property", "Value"], [
        ["Type", "ForEach"],
        ["Sequential", "No (runs in parallel)"],
        ["Batch Count", "4 (static)"],
        ["Items", "@activity('Filter_BatchSize').output.value"],
        ["Child Activity", "Execute_MigrateSingleLibrary (ExecutePipeline)"],
        ["Wait on Completion", "Yes"],
    ])

    doc.add_heading('5.6 Log_BatchComplete', level=2)
    tbl(doc, ["Property", "Value"], [
        ["Type", "SqlServerStoredProcedure"],
        ["Stored Procedure", "dbo.usp_LogBatchComplete"],
        ["Parameters", "BatchId, Status='Completed', EndTime"],
        ["Depends On", "ForEach_Library (Succeeded)"],
    ])

    doc.add_heading('6. Error Handling', level=1)
    doc.add_paragraph("If the Lookup or Filter fails, the pipeline stops and the batch is not logged as complete. "
                       "Individual library failures are handled inside PL_Migrate_Single_Library (the child pipeline). "
                       "The ForEach activity will succeed even if some child pipelines fail — the batch is marked complete "
                       "and failed libraries are retried in the next batch run.")

    doc.add_heading('7. Test Results', level=1)
    doc.add_paragraph("Run ID: 7bb1ca75-5180-4087-9870-6dea48e667a6")
    tbl(doc,
        ["Activity", "Status", "Duration"],
        [
            ["Set_BatchId", "Succeeded", "276ms"],
            ["Log_BatchStart", "Succeeded", "3,808ms"],
            ["Lookup_PendingLibraries", "Succeeded", "8,019ms"],
            ["If_NoLibrariesToProcess", "Succeeded", "978ms"],
            ["Filter_BatchSize", "Succeeded", "319ms"],
            ["Log_NoWork", "Succeeded", "254ms"],
            ["ForEach_Library", "Succeeded", "756ms"],
            ["Log_BatchComplete", "Succeeded", "3,071ms"],
        ])

    doc.save(f"{BASE}\\Pipeline_Master_Migration_Orchestrator.docx")
    print("Created: Pipeline_Master_Migration_Orchestrator.docx")

# ================================================================
# 3. PL_MIGRATE_SINGLE_LIBRARY
# ================================================================
def gen_pl_single():
    doc = make_doc()
    add_title_page(doc, "PL_Migrate_Single\nLibrary", "Pipeline Technical Documentation")

    doc.add_heading('1. Overview', level=1)
    tbl(doc, ["Property", "Value"], [
        ["Pipeline Name", "PL_Migrate_Single_Library"],
        ["ARM Template", "adf-templates/pipelines/PL_Migrate_Single_Library.json"],
        ["Purpose", "Migrates ALL files from a single SharePoint document library to ADLS Gen2"],
        ["Called By", "PL_Master_Migration_Orchestrator (ForEach_Library)"],
        ["Child Pipeline", "PL_Process_Subfolder (for each subfolder)"],
    ])

    doc.add_heading('2. Parameters', level=1)
    tbl(doc,
        ["Parameter", "Type", "Default", "Description"],
        [
            ["SiteUrl", "string", "/sites/HydroOneDocuments", "SharePoint site server-relative URL"],
            ["LibraryName", "string", "Documents", "Document library name"],
            ["ControlTableId", "int", "-", "ID of the record in MigrationControl table"],
            ["BatchId", "string", "-", "Batch identifier from parent pipeline"],
            ["ContainerName", "string", "sharepoint-migration", "ADLS destination container"],
            ["SharePointTenantUrl", "string", "https://hydroone.sharepoint.com", "SharePoint tenant base URL"],
            ["ThrottleWaitSeconds", "int", "120", "Seconds to wait when HTTP 429 throttling is detected"],
        ])

    doc.add_heading('3. Activity Flow', level=1)
    flow_diagram(doc, """
    Update_Status_InProgress
            |
            +-------------------+----------------------+
            |                                          |
    Get_RootFolderFiles (WebActivity)         Get_AllSubfolders (WebActivity)
            |                                          |
    ForEach_RootFile (parallel=10)            ForEach_Subfolder
            |                                          |
            +-- Copy_SingleFile                +-- Execute_ProcessSubfolder
            |   (SharePoint -> ADLS)           |   (calls PL_Process_Subfolder)
            |                                  |
            +-- If Copy Succeeded:             +-- Passes: SiteUrl, LibraryName,
            |   +-- Log_FileSuccess            |   FolderServerRelativeUrl,
            |   +-- Check_Throttling           |   ContainerName, SharePointTenantUrl
            |                                  |
            +-- If Copy Failed:
                +-- Log_FileFailure
            |
            +-------------------+----------------------+
            |                                          |
    Update_Status_Completed              Update_Status_Failed
    (if all succeeded)                   (if any failures)
    """)

    doc.add_heading('4. Key Activities', level=1)

    doc.add_heading('4.1 Get_RootFolderFiles', level=2)
    tbl(doc, ["Property", "Value"], [
        ["Type", "WebActivity (GET)"],
        ["Authentication", "MSI (Managed Service Identity)"],
        ["Timeout", "10 minutes"],
        ["Retry", "3 attempts, 60s interval"],
    ])
    doc.add_paragraph("URL Expression:")
    code(doc, """{SharePointTenantUrl}{SiteUrl}/_api/web/GetFolderByServerRelativeUrl(
    '{SiteUrl}/{LibraryName}'
)/Files?$select=Name,ServerRelativeUrl,Length,TimeLastModified,UniqueId&$top=5000""")
    doc.add_paragraph("Returns: JSON array of file objects with Name, ServerRelativeUrl, Length (bytes), TimeLastModified, UniqueId")

    doc.add_heading('4.2 Copy_SingleFile', level=2)
    tbl(doc, ["Property", "Value"], [
        ["Type", "Copy Activity"],
        ["Source", "DS_SharePoint_Binary_HTTP (HTTP binary download)"],
        ["Sink", "DS_ADLS_Binary_Sink (ADLS Gen2 binary write)"],
        ["Timeout", "30 minutes per file"],
        ["Retry", "3 attempts, 60s interval"],
        ["Enable Staging", "No (direct streaming)"],
    ])

    doc.add_heading('4.3 ADLS Path Mapping', level=2)
    doc.add_paragraph("Each file's destination path is computed to preserve the SharePoint folder structure:")
    tbl(doc,
        ["Parameter", "Expression", "Example"],
        [
            ["ContainerName", "@pipeline().parameters.ContainerName", "sharepoint-migration"],
            ["SiteName", "@replace(SiteUrl, '/sites/', '')", "HydroOneDocuments"],
            ["LibraryName", "@pipeline().parameters.LibraryName", "Documents"],
            ["FolderPath", "Computed by stripping site/library prefix from server-relative URL", "Reports/2024/Q1"],
            ["FileName", "@item().Name", "Q1_Report.pdf"],
        ])
    doc.add_paragraph("Result: sharepoint-migration/HydroOneDocuments/Documents/Reports/2024/Q1/Q1_Report.pdf")

    doc.add_heading('4.4 Check_Throttling', level=2)
    tbl(doc, ["Property", "Value"], [
        ["Type", "IfCondition"],
        ["Expression", "@contains(string(activity('Copy_SingleFile').output), '429')"],
        ["If True", "Wait activity pauses for ThrottleWaitSeconds (default: 120s)"],
        ["If False", "Continue to next file"],
    ])

    doc.add_heading('4.5 Log_FileSuccess / Log_FileFailure', level=2)
    tbl(doc, ["Property", "Value"], [
        ["Type", "SqlServerStoredProcedure"],
        ["Stored Procedure", "dbo.usp_LogFileAudit"],
        ["Key Parameters", "FileName, SourcePath, DestinationPath, FileSizeBytes, MigrationStatus, PipelineRunId, ErrorDetails"],
    ])

    doc.add_heading('5. Error Handling', level=1)
    tbl(doc,
        ["Error", "Pipeline Behavior"],
        [
            ["HTTP 429 (Throttled)", "Wait ThrottleWaitSeconds, then retry (up to 3 times)"],
            ["HTTP 401/403", "Log error, skip file, continue with remaining files"],
            ["HTTP 404", "Log as 'file not found at source', skip, continue"],
            ["Timeout", "Retry with same timeout (30 min per file)"],
            ["All files succeed", "Library status updated to Completed"],
            ["Any files fail", "Library status updated to Failed, RetryCount incremented"],
        ])

    doc.save(f"{BASE}\\Pipeline_Migrate_Single_Library.docx")
    print("Created: Pipeline_Migrate_Single_Library.docx")

# ================================================================
# 4. PL_PROCESS_SUBFOLDER
# ================================================================
def gen_pl_subfolder():
    doc = make_doc()
    add_title_page(doc, "PL_Process_Subfolder", "Pipeline Technical Documentation")

    doc.add_heading('1. Overview', level=1)
    tbl(doc, ["Property", "Value"], [
        ["Pipeline Name", "PL_Process_Subfolder"],
        ["ARM Template", "adf-templates/pipelines/PL_Process_Subfolder.json"],
        ["Purpose", "Processes all files within a single SharePoint subfolder"],
        ["Called By", "PL_Migrate_Single_Library (ForEach_Subfolder)"],
        ["Children", "None (leaf pipeline)"],
    ])
    note(doc, "This pipeline does NOT recurse into sub-subfolders. ADF does not support recursive pipeline self-references. Subfolders are flattened by the parent pipeline's enumeration.")

    doc.add_heading('2. Parameters', level=1)
    tbl(doc,
        ["Parameter", "Type", "Description"],
        [
            ["SiteUrl", "string", "SharePoint site server-relative URL (e.g., /sites/HydroOne)"],
            ["LibraryName", "string", "Document library name (e.g., Documents)"],
            ["FolderServerRelativeUrl", "string", "Full server-relative URL of the folder to process"],
            ["ContainerName", "string", "ADLS container name"],
            ["SharePointTenantUrl", "string", "SharePoint tenant base URL"],
        ])

    doc.add_heading('3. Activity Flow', level=1)
    flow_diagram(doc, """
    Get_FolderFiles (WebActivity GET)
            |
            | Returns: array of files in this folder
            |
    ForEach_File (parallel=10)
            |
            +-- Copy_File (Binary: SharePoint HTTP -> ADLS Gen2)
            |
            +-- If Succeeded:
            |       +-- Log_Success (usp_LogFileAudit, Status='Success')
            |
            +-- If Failed:
                    +-- Log_Failure (usp_LogFileAudit, Status='Failed')
    """)

    doc.add_heading('4. SharePoint API Call', level=1)
    doc.add_paragraph("The Get_FolderFiles WebActivity calls:")
    code(doc, """{SharePointTenantUrl}{SiteUrl}/_api/web/GetFolderByServerRelativeUrl(
    '{FolderServerRelativeUrl}'
)/Files?$select=Name,ServerRelativeUrl,Length,TimeLastModified,UniqueId&$top=5000""")
    doc.add_paragraph("Authentication: MSI (Managed Service Identity)")

    doc.add_heading('5. ADLS Path Computation', level=1)
    doc.add_paragraph("The folder path in ADLS is computed by stripping the site and library prefix from the SharePoint server-relative URL:")
    code(doc, """FolderPath = replace(
    replace(
        FolderServerRelativeUrl,
        concat(SiteUrl, '/', LibraryName, '/'),
        ''
    ),
    concat(SiteUrl, '/', LibraryName),
    ''
)""")
    doc.add_paragraph("Example:")
    tbl(doc,
        ["SharePoint Path", "ADLS Path"],
        [
            ["/sites/HydroOne/Documents/Reports/2024", "sharepoint-migration/HydroOne/Documents/Reports/2024/"],
            ["/sites/HydroOne/Documents/Policies", "sharepoint-migration/HydroOne/Documents/Policies/"],
        ])

    doc.add_heading('6. Design Decision: No Recursion', level=1)
    doc.add_paragraph(
        "ADF does not allow a pipeline to call itself (circular reference error). The original design "
        "attempted recursive subfolder processing, but this was removed. Instead, the parent pipeline "
        "(PL_Migrate_Single_Library) enumerates ALL subfolders at all levels and passes each one to this "
        "pipeline as a flat list. This approach is simpler and avoids ADF's recursion limitation.")

    doc.save(f"{BASE}\\Pipeline_Process_Subfolder.docx")
    print("Created: Pipeline_Process_Subfolder.docx")

# ================================================================
# 5. PL_VALIDATION
# ================================================================
def gen_pl_validation():
    doc = make_doc()
    add_title_page(doc, "PL_Validation", "Pipeline Technical Documentation")

    doc.add_heading('1. Overview', level=1)
    tbl(doc, ["Property", "Value"], [
        ["Pipeline Name", "PL_Validation"],
        ["ARM Template", "adf-templates/pipelines/PL_Validation.json"],
        ["Purpose", "Post-migration validation comparing expected file counts (control table) with actual migration results (audit log)"],
        ["Trigger", "Manual (run after migration batches complete)"],
        ["Test Status", "PASSED (Run ID: 92158504, all 8 activities succeeded, library marked Validated)"],
    ])

    doc.add_heading('1.1 Validation Approach', level=2)
    doc.add_paragraph(
        "This pipeline uses SQL-only validation — it does NOT call SharePoint during validation. "
        "It compares the expected file count from the MigrationControl table with the actual count of "
        "successful entries in the MigrationAuditLog table. This design ensures validation can run "
        "without consuming SharePoint API quota.")

    doc.add_heading('2. Parameters', level=1)
    tbl(doc,
        ["Parameter", "Type", "Default", "Description"],
        [
            ["SharePointTenantUrl", "string", "https://hydroone.sharepoint.com", "Reserved for future SharePoint-direct validation"],
            ["ValidateAll", "bool", "false", "true = re-validate all completed libraries; false = only validate Pending/NULL"],
        ])
    note(doc, "The ValidateAll boolean is converted to 1/0 in the SQL query using @{if(pipeline().parameters.ValidateAll, 1, 0)} because SQL does not understand ADF's True/False rendering.")

    doc.add_heading('3. Activity Flow', level=1)
    flow_diagram(doc, """
    Lookup_CompletedLibraries
            |
            | Returns: libraries with Status='Completed' and ValidationStatus pending
            |
    ForEach_ValidateLibrary (parallel=5)
            |
            +-- Lookup_DestinationFileCount
            |       SQL: COUNT audit log entries for this library
            |       Returns: ActualFileCount, ActualSizeBytes, SuccessCount, FailedCount
            |
            +-- Compare_And_Log (stored procedure: usp_LogValidationResult)
            |       Updates MigrationControl with actual counts
            |
            +-- If_Discrepancy
                    |
                    +-- Condition: FailedCount > 0 OR ExpectedFileCount != SuccessCount
                    |
                    +-- TRUE:  Flag_Discrepancy
                    |          Calls usp_UpdateValidationStatus('Discrepancy', details)
                    |
                    +-- FALSE: Mark_Validated
                               Calls usp_UpdateValidationStatus('Validated', NULL)
            |
    Generate_ValidationReport
            |
            | SQL: SELECT ValidatedCount, DiscrepancyCount, PendingCount, TotalLibraries
            |
    Set_ValidationSummary
            |
            | Variable: "Validation Complete - Validated: X, Discrepancies: Y, Pending: Z"
    """)

    doc.add_heading('4. Activity Details', level=1)

    doc.add_heading('4.1 Lookup_CompletedLibraries', level=2)
    doc.add_paragraph("SQL Query:")
    code(doc, """SELECT Id, SiteUrl, LibraryName,
    FileCount AS ExpectedFileCount,
    TotalSizeBytes AS ExpectedSizeBytes
FROM dbo.MigrationControl
WHERE Status = 'Completed'
AND (@{if(pipeline().parameters.ValidateAll, 1, 0)} = 1
    OR ValidationStatus IS NULL
    OR ValidationStatus = 'Pending')""")

    doc.add_heading('4.2 Lookup_DestinationFileCount', level=2)
    doc.add_paragraph("SQL Query:")
    code(doc, """SELECT
    COUNT(*) AS ActualFileCount,
    SUM(FileSizeBytes) AS ActualSizeBytes,
    SUM(CASE WHEN MigrationStatus = 'Success' THEN 1 ELSE 0 END) AS SuccessCount,
    SUM(CASE WHEN MigrationStatus = 'Failed' THEN 1 ELSE 0 END) AS FailedCount
FROM dbo.MigrationAuditLog
WHERE SourcePath LIKE '{SiteUrl}/{LibraryName}%'""")

    doc.add_heading('4.3 If_Discrepancy Expression', level=2)
    code(doc, "@or(greater(FailedCount, 0), not(equals(item().ExpectedFileCount, SuccessCount)))")
    doc.add_paragraph("This triggers a discrepancy flag if:")
    doc.add_paragraph("- Any files have MigrationStatus = 'Failed', OR", style='List Bullet')
    doc.add_paragraph("- The expected file count does not match the successful migration count", style='List Bullet')

    doc.add_heading('5. Validation Outcomes', level=1)
    tbl(doc,
        ["Outcome", "Condition", "Action Taken"],
        [
            ["Validated", "ExpectedFileCount == SuccessCount AND FailedCount == 0", "ValidationStatus set to 'Validated'"],
            ["Discrepancy", "ExpectedFileCount != SuccessCount OR FailedCount > 0", "ValidationStatus set to 'Discrepancy' with detail string"],
        ])
    doc.add_paragraph("Discrepancy detail format: 'Expected: 100, Actual Success: 98, Failed: 2'")

    doc.add_heading('6. Test Results', level=1)
    doc.add_paragraph("Run ID: 92158504-c822-46a3-9a4f-2cbd780986f6")
    tbl(doc,
        ["Activity", "Status", "Duration"],
        [
            ["Lookup_CompletedLibraries", "Succeeded", "26,107ms"],
            ["ForEach_ValidateLibrary", "Succeeded", "40,775ms"],
            ["Lookup_DestinationFileCount", "Succeeded", "24,090ms"],
            ["Compare_And_Log", "Succeeded", "3,717ms"],
            ["If_Discrepancy", "Succeeded", "5,365ms"],
            ["Mark_Validated", "Succeeded", "2,817ms"],
            ["Generate_ValidationReport", "Succeeded", "13,051ms"],
            ["Set_ValidationSummary", "Succeeded", "276ms"],
        ])
    doc.add_paragraph("Result: 10 expected files, 10 migrated, 0 failed. Library marked as Validated.")

    doc.save(f"{BASE}\\Pipeline_Validation.docx")
    print("Created: Pipeline_Validation.docx")

# ================================================================
# 6. PL_INCREMENTAL_SYNC
# ================================================================
def gen_pl_incremental():
    doc = make_doc()
    add_title_page(doc, "PL_Incremental_Sync", "Pipeline Technical Documentation")

    doc.add_heading('1. Overview', level=1)
    tbl(doc, ["Property", "Value"], [
        ["Pipeline Name", "PL_Incremental_Sync"],
        ["ARM Template", "adf-templates/pipelines/PL_Incremental_Sync.json"],
        ["Purpose", "Delta/incremental synchronization for ongoing sync after initial migration"],
        ["Trigger", "Tumbling window (every 6 hours) or manual"],
        ["Prerequisite", "Libraries must have Status='Completed', ValidationStatus='Validated', EnableIncrementalSync=1"],
    ])

    doc.add_heading('1.1 How It Works', level=2)
    doc.add_paragraph(
        "This pipeline queries SharePoint for files modified after a stored watermark timestamp. "
        "Only changed or new files are copied to ADLS, overwriting any existing version. After "
        "successful sync, the watermark is updated to the current time.")

    doc.add_heading('2. Parameters', level=1)
    tbl(doc,
        ["Parameter", "Type", "Default", "Description"],
        [
            ["SharePointTenantUrl", "string", "https://hydroone.sharepoint.com", "SharePoint tenant URL"],
            ["ContainerName", "string", "sharepoint-migration", "ADLS container name"],
        ])

    doc.add_heading('3. Activity Flow', level=1)
    flow_diagram(doc, """
    Lookup_LibrariesForSync
            |
            | SQL: SELECT libraries WHERE Status='Completed'
            |      AND ValidationStatus='Validated' AND EnableIncrementalSync=1
            |      LEFT JOIN IncrementalWatermark to get LastModifiedDate
            |
    ForEach_LibrarySync (parallel=4)
            |
            +-- Get_ModifiedFiles (WebActivity GET)
            |       SharePoint REST API with Modified date filter
            |       URL: /_api/web/lists/getbytitle('{lib}')/items
            |            ?$filter=Modified ge datetime'{lastModifiedDate}'
            |
            +-- If_HasModifiedFiles
                    |
                    +-- Condition: length(results) > 0
                    |
                    +-- TRUE:
                    |       +-- ForEach_ModifiedFile (parallel=10)
                    |       |       +-- Copy_ModifiedFile (SharePoint -> ADLS)
                    |       |       +-- Log_IncrementalCopy (usp_LogFileAudit)
                    |       |
                    |       +-- Update_Watermark (usp_UpdateWatermark)
                    |               Sets LastModifiedDate = utcNow()
                    |               Sets LastSyncTime = utcNow()
                    |
                    +-- FALSE: No action (no files changed since last sync)
            |
    Log_SyncComplete (usp_LogSyncRun)
            |
            | Logs: PipelineRunId, SyncType='Incremental',
            |       LibrariesProcessed, CompletionTime
    """)

    doc.add_heading('4. SharePoint Modified Files Query', level=1)
    doc.add_paragraph("The Get_ModifiedFiles WebActivity calls:")
    code(doc, """{SharePointTenantUrl}{SiteUrl}/_api/web/lists/getbytitle('{LibraryName}')/items
    ?$filter=Modified ge datetime'{LastModifiedDate}'
    &$select=FileRef,FileLeafRef,File_x0020_Size,Modified,UniqueId
    &$expand=File
    &$top=5000""")
    doc.add_paragraph("This returns only files modified after the watermark date. The $top=5000 limit handles most libraries. For libraries with more than 5,000 changes between syncs, the query may need pagination.")

    doc.add_heading('5. Watermark Management', level=1)
    tbl(doc,
        ["Table", "Column", "Description"],
        [
            ["IncrementalWatermark", "SiteUrl", "SharePoint site URL"],
            ["IncrementalWatermark", "LibraryName", "Document library name"],
            ["IncrementalWatermark", "LastModifiedDate", "High watermark — only files modified after this date are synced"],
            ["IncrementalWatermark", "LastSyncTime", "When the last successful sync completed"],
        ])
    doc.add_paragraph()
    doc.add_paragraph("The stored procedure usp_UpdateWatermark uses MERGE to either update an existing watermark or insert a new one:")
    code(doc, """MERGE dbo.IncrementalWatermark AS target
USING (SELECT @SiteUrl, @LibraryName) AS source
ON target.SiteUrl = source.SiteUrl AND target.LibraryName = source.LibraryName
WHEN MATCHED THEN UPDATE SET LastModifiedDate = @LastModifiedDate, LastSyncTime = @LastSyncTime
WHEN NOT MATCHED THEN INSERT (...) VALUES (...)""")

    doc.add_heading('6. Initial Watermark Setup', level=1)
    doc.add_paragraph("Before starting incremental sync, initialize the watermark table:")
    code(doc, """INSERT INTO dbo.IncrementalWatermark (SiteUrl, LibraryName, LastModifiedDate, LastSyncTime)
SELECT SiteUrl, LibraryName, EndTime, GETUTCDATE()
FROM dbo.MigrationControl
WHERE EnableIncrementalSync = 1;""")
    note(doc, "The initial watermark is set to the migration EndTime so that only files modified AFTER the initial migration are synced.")

    doc.add_heading('7. Sync Frequency', level=1)
    tbl(doc,
        ["Frequency", "Use Case", "API Impact"],
        [
            ["Every 6 hours", "Default — balances freshness with API usage", "Low"],
            ["Every 1 hour", "High-change environments", "Medium"],
            ["Every 15 minutes", "Near-real-time requirements", "High (watch for throttling)"],
            ["Daily", "Low-change environments", "Minimal"],
        ])

    doc.save(f"{BASE}\\Pipeline_Incremental_Sync.docx")
    print("Created: Pipeline_Incremental_Sync.docx")

# ================================================================
# GENERATE ALL
# ================================================================
if __name__ == "__main__":
    gen_architecture()
    gen_pl_master()
    gen_pl_single()
    gen_pl_subfolder()
    gen_pl_validation()
    gen_pl_incremental()
    print("\nAll 6 Word documents generated successfully!")
